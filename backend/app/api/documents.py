import os
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from typing import Optional

from app.database import get_db
from app.core.deps import get_current_user
from app.core.config import settings
from app.core.project_access import require_project_write, require_project_read, require_model_access
from app.models.user import User, UserRole
from app.models.project import Project
from app.models.document import Document, DocumentStatus
from app.models.model import SysModel, ModelElement
from app.models.template import Template
from app.schemas.document import DocumentGenerate, DocumentOut, DocumentListResponse
from app.services.log_service import create_log
from app.models.log import LogType, ResultStatus
from jinja2.sandbox import SandboxedEnvironment

router = APIRouter(prefix="/api/v1/documents", tags=["documents"])

# ── HTML → DOCX 转换 ──────────────────────────

def _add_inline_content(paragraph, element):
    """将 HTML 内联元素（strong/em/code/br 等）转为 python-docx runs。"""
    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text.strip():
                paragraph.add_run(text)
        elif isinstance(child, Tag):
            txt = child.get_text()
            if not txt.strip():
                continue
            run = paragraph.add_run(txt)
            if child.name in ('strong', 'b'):
                run.bold = True
            elif child.name in ('em', 'i'):
                run.italic = True
            elif child.name == 'code':
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            elif child.name == 'br':
                run.add_break()


def _set_cell_text(cell, td_tag):
    """填充表格单元格内容，保留加粗/斜体等格式。"""
    for child in td_tag.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text.strip():
                p = cell.paragraphs[0]
                p.add_run(text)
        elif isinstance(child, Tag):
            if child.name == 'p':
                _add_inline_content(cell.add_paragraph(), child)
            else:
                _add_inline_content(cell.paragraphs[0], child)


def html_to_docx(html: str, docx_doc):
    """将 HTML 内容解析并写入 python-docx Document 对象。"""
    soup = BeautifulSoup(html, 'html.parser')
    body = soup.find('body') or soup

    for el in body.children:
        if not isinstance(el, Tag):
            continue
        tag = el.name

        # 跳过样式/元数据
        if tag in ('style', 'head', 'meta', 'script', 'title', 'link'):
            continue

        # 标题
        if tag in ('h1', 'h2', 'h3', 'h4'):
            level = int(tag[1])
            docx_doc.add_heading(el.get_text(strip=True), level=level)
            continue

        # 表格
        if tag == 'table':
            rows = el.find_all('tr')
            if not rows:
                continue
            cols = max(len(r.find_all(['th', 'td'])) for r in rows)
            table = docx_doc.add_table(rows=len(rows), cols=cols, style='Table Grid')
            for i, tr_tag in enumerate(rows):
                cells = tr_tag.find_all(['th', 'td'])
                for j, td_tag in enumerate(cells):
                    if j >= cols:
                        break
                    _set_cell_text(table.rows[i].cells[j], td_tag)
            continue

        # 段落
        if tag == 'p':
            p = docx_doc.add_paragraph()
            _add_inline_content(p, el)
            continue

        # 列表
        if tag in ('ul', 'ol'):
            for li in el.find_all('li', recursive=False):
                p = docx_doc.add_paragraph(style='List Bullet')
                _add_inline_content(p, li)
            continue

        # 代码块
        if tag == 'pre':
            p = docx_doc.add_paragraph()
            run = p.add_run(el.get_text())
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            continue

        # div/section/article 及其它容器 → 递归处理子元素
        if tag in ('div', 'section', 'article', 'main', 'header', 'footer', 'span'):
            for child in el.children:
                if isinstance(child, Tag):
                    # 递归调用自身，用子元素的 HTML 创建临时 doc
                    html_to_docx(str(child), docx_doc)
                elif isinstance(child, NavigableString):
                    txt = str(child).strip()
                    if txt:
                        docx_doc.add_paragraph(txt)
            continue

        # 其它块级元素——按段落处理
        txt = el.get_text(strip=True)
        if txt:
            docx_doc.add_paragraph(txt)


# ── HTML → PDF 中文支持 ────────────────────────

def _register_cjk_font() -> str | None:
    """注册系统 CJK 字体到 ReportLab，返回字体名称，失败返回 None。"""
    import os.path as osp
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    font_candidates = [
        ("C:/Windows/Fonts/simhei.ttf", "SimHei"),
        ("C:/Windows/Fonts/msyh.ttf", "Microsoft YaHei"),
        ("C:/Windows/Fonts/msyh.ttc", "Microsoft YaHei"),
        ("C:/Windows/Fonts/simsun.ttc", "SimSun"),
        ("C:/Windows/Fonts/simkai.ttf", "KaiTi"),
        ("C:/Windows/Fonts/fangsong.ttf", "FangSong"),
    ]
    for path, name in font_candidates:
        if osp.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(name, path))
                return name
            except Exception:
                continue
    return None


# 模块加载时注册中文字体
_CJK_FONT_NAME = _register_cjk_font()


def _prepare_html_for_pdf(html: str) -> str:
    """注入字体 CSS 到 HTML 中供 xhtml2pdf 使用。"""
    if not _CJK_FONT_NAME:
        return html
    # 将字体注册到 xhtml2pdf 的字体映射表，否则 xhtml2pdf 不认
    from xhtml2pdf import default
    font_key = _CJK_FONT_NAME.lower()
    if font_key not in default.DEFAULT_FONT:
        default.DEFAULT_FONT[font_key] = _CJK_FONT_NAME
        # 覆盖 html 默认字体（原为 Helvetica）
        default.DEFAULT_CSS = default.DEFAULT_CSS.replace(
            "font-family: Helvetica;",
            f"font-family: {_CJK_FONT_NAME};",
        )
    font_css = f'<style>html {{ font-family: "{_CJK_FONT_NAME}", sans-serif; }}</style>\n'
    if "</head>" in html:
        return html.replace("</head>", font_css + "</head>")
    elif "<body" in html:
        idx = html.find("<body")
        idx = html.find(">", idx) + 1
        return html[:idx] + font_css + html[idx:]
    else:
        return font_css + html

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from weasyprint import HTML
    HAS_WEASYPRINT = True
except Exception:
    HAS_WEASYPRINT = False

try:
    from fpdf import FPDF
    HAS_FPDF = True
except ImportError:
    HAS_FPDF = False

try:
    from bs4 import BeautifulSoup, Tag, NavigableString
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

try:
    from xhtml2pdf import pisa
    HAS_XHTML2PDF = True
except ImportError:
    HAS_XHTML2PDF = False


@router.get("", response_model=DocumentListResponse)
def list_documents(
    project_id: Optional[int] = Query(None, description="按项目筛选"),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    q = db.query(Document)
    if project_id is not None:
        require_project_read(db, current_user, project_id)
        q = q.filter(Document.project_id == project_id)
    elif current_user.role == UserRole.MEMBER.value:
        q = q.join(Project, Project.id == Document.project_id).filter(Project.owner_id == current_user.id)
    total = q.count()
    rows = (
        q.order_by(Document.created_at.desc())
        .offset((page - 1) * page_size)
        .limit(page_size)
        .all()
    )
    return DocumentListResponse(items=rows, total=total, page=page, page_size=page_size)


@router.post("/generate", response_model=DocumentOut)
def generate_document(doc_in: DocumentGenerate, db: Session = Depends(get_db),
                      current_user: User = Depends(get_current_user)):
    require_project_write(db, current_user, doc_in.project_id)
    model = db.query(SysModel).filter(SysModel.id == doc_in.model_id).first()
    if not model:
        raise HTTPException(status_code=404, detail="Model not found")
    if model.project_id != doc_in.project_id:
        raise HTTPException(status_code=400, detail="模型不属于所选项目")
    require_model_access(db, current_user, model, write=False)
    template = db.query(Template).filter(Template.id == doc_in.template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    elements = db.query(ModelElement).filter(ModelElement.model_id == model.id).all()

    def get_element_by_id(element_id: str):
        row = (
            db.query(ModelElement)
            .filter(ModelElement.model_id == model.id, ModelElement.element_id == element_id)
            .first()
        )
        if not row:
            return None
        return {
            "element_id": row.element_id,
            "name": row.element_name,
            "type": row.element_type,
            "description": row.description,
        }

    def get_requirements_by_module(module_name: str):
        return [
            {"name": e.element_name, "type": e.element_type, "description": e.description}
            for e in elements
            if e.element_type == "Requirement" and module_name in (e.element_name or "")
        ]

    def get_blocks_by_package(package_name: str):
        return [
            {"name": e.element_name, "type": e.element_type, "description": e.description}
            for e in elements
            if e.element_type == "Block" and package_name in (e.element_name or "")
        ]

    # Build context
    context = {
        "project_id": doc_in.project_id,
        "model_name": model.name,
        "model_version": model.version_tag,
        "template_name": template.name,
        "generate_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "elements": [{
            "name": e.element_name,
            "type": e.element_type,
            "description": e.description,
        } for e in elements],
        "get_element_by_id": get_element_by_id,
        "get_requirements_by_module": get_requirements_by_module,
        "get_blocks_by_package": get_blocks_by_package,
    }

    # Render（沙箱环境，防止模板执行危险代码）
    try:
        env = SandboxedEnvironment(autoescape=True)
        jinja_tpl = env.from_string(template.content)
        rendered = jinja_tpl.render(**context)
    except Exception as e:
        doc = Document(
            project_id=doc_in.project_id, model_id=doc_in.model_id,
            template_id=doc_in.template_id,
            document_name=f"{model.name}_{template.name}",
            status=DocumentStatus.FAILED.value,
            generate_message=str(e),
            operator_id=current_user.id,
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)
        create_log(db, LogType.OPERATION, current_user.id, "document",
                   f"Generate document failed: {e}", ResultStatus.FAILED)
        return doc

    doc = Document(
        project_id=doc_in.project_id, model_id=doc_in.model_id,
        template_id=doc_in.template_id,
        document_name=f"{model.name}_{template.name}",
        content=rendered,
        status=DocumentStatus.SUCCESS.value,
        operator_id=current_user.id,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    create_log(db, LogType.OPERATION, current_user.id, "document",
               f"Generated document {doc.document_name}")
    return doc


@router.get("/{doc_id}", response_model=DocumentOut)
def get_document(doc_id: int, db: Session = Depends(get_db),
                 current_user: User = Depends(get_current_user)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    require_project_read(db, current_user, doc.project_id)
    return doc


@router.get("/{doc_id}/preview")
def preview_document(doc_id: int, db: Session = Depends(get_db),
                     current_user: User = Depends(get_current_user)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    require_project_read(db, current_user, doc.project_id)
    return {"content": doc.content, "document_name": doc.document_name}


@router.get("/{doc_id}/download")
def download_document(doc_id: int, fmt: str = "docx",
                      db: Session = Depends(get_db),
                      current_user: User = Depends(get_current_user)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    require_project_write(db, current_user, doc.project_id)
    if fmt == "docx" and HAS_DOCX and HAS_BS4:
        os.makedirs(settings.document_dir, exist_ok=True)
        filepath = os.path.join(settings.document_dir, f"{doc.document_name}_{doc.id}.docx")
        d = DocxDocument()
        d.add_heading(doc.document_name, 0)
        html_to_docx(doc.content, d)
        d.save(filepath)
        doc.file_path = filepath
        doc.export_format = "docx"
        db.commit()
        return FileResponse(filepath, filename=f"{doc.document_name}.docx")
    elif fmt == "pdf" and (HAS_WEASYPRINT or HAS_XHTML2PDF or HAS_FPDF):
        os.makedirs(settings.document_dir, exist_ok=True)
        filepath = os.path.join(settings.document_dir, f"{doc.document_name}_{doc.id}.pdf")
        if HAS_WEASYPRINT:
            HTML(string=doc.content).write_pdf(filepath)
        elif HAS_XHTML2PDF:
            pdf_content = _prepare_html_for_pdf(doc.content)
            with open(filepath, "wb") as f:
                pisa.CreatePDF(pdf_content, dest=f, encoding="utf-8")
        else:
            import re
            clean = re.sub(r'<[^>]+>', '', doc.content)
            lines = clean.split('\n')
            pdf = FPDF()
            pdf.add_page()
            import os.path as osp
            font_paths = [
                "C:/Windows/Fonts/msyh.ttc",
                "C:/Windows/Fonts/simsun.ttc",
                "C:/Windows/Fonts/msyh.ttf",
                "C:/Windows/Fonts/simhei.ttf",
            ]
            cjk_font = None
            for fp in font_paths:
                if osp.exists(fp):
                    cjk_font = fp
                    break
            if cjk_font:
                pdf.add_font("CJK", "", cjk_font, uni=True)
            line_height = 7
            for line in lines:
                pdf.set_x(10)
                line = line.strip()
                if not line:
                    continue
                if cjk_font:
                    pdf.set_font("CJK", size=12)
                else:
                    pdf.set_font("Helvetica", size=12)
                size = 12
                if line.startswith('# '):
                    size = 18
                    text = line[2:]
                elif line.startswith('## '):
                    size = 15
                    text = line[3:]
                else:
                    text = line
                pdf.set_font_size(size)
                try:
                    pdf.multi_cell(0, line_height, text=text)
                except Exception:
                    pdf.set_x(10)
                    pdf.set_font_size(10)
                    safe = ''.join(c if ord(c) < 256 else '?' for c in text)
                    pdf.multi_cell(0, line_height, text=safe)
            pdf.output(filepath)
        doc.file_path = filepath
        doc.export_format = "pdf"
        db.commit()
        return FileResponse(filepath, filename=f"{doc.document_name}.pdf", media_type="application/pdf")
    elif fmt == "html":
        return {"content": doc.content}
    else:
        if fmt == "pdf" and not HAS_WEASYPRINT and not HAS_XHTML2PDF and not HAS_FPDF:
            raise HTTPException(status_code=503, detail="PDF 导出需要 WeasyPrint、xhtml2pdf 或 fpdf2，当前环境未安装")
        raise HTTPException(status_code=400, detail=f"Unsupported format: {fmt}")


@router.delete("/{doc_id}")
def delete_document(doc_id: int, db: Session = Depends(get_db),
                    current_user: User = Depends(get_current_user)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    require_project_write(db, current_user, doc.project_id)
    if doc.file_path and os.path.exists(doc.file_path):
        os.remove(doc.file_path)
    db.delete(doc)
    db.commit()
    return {"detail": "Document deleted"}
