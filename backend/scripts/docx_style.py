# -*- coding: utf-8 -*-
"""Word 文档样式工具（正文：宋体五号、单倍行距）"""
from __future__ import annotations

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# 五号字 = 10.5pt；正文单倍行距
BODY_FONT_PT = 10.5
BODY_LINE_SPACING = 1.0


def set_east_asia_font(run, name: str = "宋体") -> None:
    run.font.name = name
    r = run._element.get_or_add_rPr()
    rfonts = r.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)


def _apply_body_paragraph_format(pf) -> None:
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = BODY_LINE_SPACING


def set_cell_shading(cell, fill: str = "D9D9D9") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def setup_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.5)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(BODY_FONT_PT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    pf = normal.paragraph_format
    _apply_body_paragraph_format(pf)
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def setup_document_compact(doc: Document) -> None:
    """合并正文：五号、单倍行距，段后 3 磅便于阅读。"""
    setup_document(doc)
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)


def add_run(paragraph, text: str, *, size: float = BODY_FONT_PT, bold: bool = False, font: str = "宋体"):
    run = paragraph.add_run(text)
    set_east_asia_font(run, font)
    run.font.size = Pt(size)
    run.font.bold = bold
    return run


def add_para(
    doc: Document,
    text: str,
    *,
    size: float = BODY_FONT_PT,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    indent_cm: float = 0,
    space_before: float = 0,
    space_after: float = 0,
):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.left_indent = Cm(indent_cm)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    _apply_body_paragraph_format(pf)
    add_run(p, text, size=size, bold=bold)
    return p


def add_heading(
    doc: Document,
    text: str,
    level: int = 1,
    *,
    page_break_before: bool = False,
    compact: bool = False,
):
    """章节标题：五号加粗（与正文同字号）。"""
    bolds = {1: True, 2: True, 3: False}
    indents = {1: 0, 2: 0.75, 3: 1.5}
    sb = 12 if level == 1 and not compact else (6 if not compact else 3)
    p = add_para(
        doc,
        text,
        size=BODY_FONT_PT,
        bold=bolds.get(level, False),
        indent_cm=indents.get(level, 0),
        space_before=sb,
        space_after=0,
    )
    if page_break_before:
        p.paragraph_format.page_break_before = True
    return p


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)


def _set_cell_body_format(p) -> None:
    _apply_body_paragraph_format(p.paragraph_format)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    col_widths_cm: list[float] | None = None,
    header_fill: bool = True,
    font_size: float = BODY_FONT_PT,
    align_center: bool = True,
    spacing_after: bool = True,
):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_body_format(p)
        add_run(p, h, size=font_size, bold=True)
        if header_fill:
            set_cell_shading(hdr_cells[i])
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
            _set_cell_body_format(p)
            add_run(p, str(val), size=font_size)
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    if spacing_after:
        gap = doc.add_paragraph()
        gap.paragraph_format.space_before = Pt(0)
        gap.paragraph_format.space_after = Pt(0)
    return table


def add_cover_block(
    doc: Document,
    doc_title: str,
    doc_id: str,
    subtitle: str,
    page_total: int,
    page_no: int,
):
    """模板封面签署表（封面表内仍可用略大字号）"""
    t = doc.add_table(rows=7, cols=6)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)

    def cell(r, c, text, *, bold=False, size=BODY_FONT_PT, center=True):
        cell_obj = t.rows[r].cells[c]
        cell_obj.text = ""
        p = cell_obj.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        lines = str(text).replace("\x0b", "\n").split("\n")
        for i, line in enumerate(lines):
            if i > 0:
                p.add_run().add_break()
            add_run(p, line, size=size, bold=bold)

    cell(0, 0, doc_title, bold=True, size=12)
    cell(0, 5, doc_id, size=BODY_FONT_PT)
    cell(1, 0, "标记")
    cell(1, 1, "数量")
    cell(1, 2, "修改单号")
    cell(1, 3, "签字")
    cell(1, 4, "日期")
    cell(2, 0, "编制")
    cell(2, 1, "1")
    cell(2, 2, "—")
    cell(2, 3, "吴一昊")
    cell(2, 4, "2026-05-20")
    cell(2, 5, "会签")
    cell(3, 3, "申博文")
    cell(3, 4, "2026-05-21")
    cell(3, 5, subtitle, size=BODY_FONT_PT)
    cell(4, 0, "校对")
    cell(4, 1, "1")
    cell(4, 2, "—")
    cell(4, 3, "林文诚")
    cell(4, 4, "2026-05-22")
    cell(4, 5, f"共 {page_total} 页")
    cell(5, 0, "审核")
    cell(5, 1, "1")
    cell(5, 2, "—")
    cell(5, 3, "张德平")
    cell(5, 4, "2026-05-23")
    cell(5, 5, f"第 {page_no} 页")
    cell(6, 0, "批准")
    cell(6, 1, "1")
    cell(6, 2, "—")
    cell(6, 3, "张德平")
    cell(6, 4, "2026-05-24")
    cell(6, 5, "会签")
    doc.add_paragraph()


def add_revision_table(doc: Document):
    add_para(doc, "修订记录", size=BODY_FONT_PT, bold=True, space_after=6)
    add_table(
        doc,
        ["版本号", "修订状态", "简要说明修订内容和范围", "修订日期", "修订人", "批准日期"],
        [["V1.0", "A", "首版发布", "2026-05-24", "吴一昊", "2026-05-24（张德平）"]],
        col_widths_cm=[1.5, 1.5, 6.5, 2.0, 2.0, 2.5],
        font_size=BODY_FONT_PT,
    )
    add_para(doc, "注：修订状态栏填写：A—增加，M—修改，D—删除。", size=BODY_FONT_PT)


def add_toc(doc: Document, items: list[str]):
    add_para(doc, "目　录", size=BODY_FONT_PT, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    for item in items:
        add_para(doc, item, size=BODY_FONT_PT, space_after=0)


def add_toc_with_pages(doc: Document, entries: list[tuple[int, str, int]]):
    """目录：居中标题 + TOC 条目（点线页码右对齐 15.35cm）"""
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

    add_para(
        doc,
        "目  录",
        size=BODY_FONT_PT,
        bold=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=12,
        space_after=12,
    )
    tab_pos = Cm(15.35)
    for level, title, page in entries:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Cm(0)
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        _apply_body_paragraph_format(pf)
        pf.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        add_run(p, title, size=BODY_FONT_PT, bold=(level == 0))
        add_run(p, "\t" + str(page), size=BODY_FONT_PT)


def add_standard_front_matter(
    doc: Document,
    toc_entries: list[tuple[int, str, int]],
    *,
    page_total: int = 40,
):
    """合并文档首页：封面签署表（含标题）+ 修订记录 + 合并目录。"""
    add_cover_block(
        doc,
        "基于 SysML 模型的文档自动生成系统\n软件用户与测试文档 V1.0",
        "SysMLDocGen-DOC-MERGE-V1.0",
        "软件用户与测试文档",
        page_total,
        1,
    )
    add_revision_table(doc)
    add_toc_with_pages(doc, toc_entries)
    doc.add_page_break()
