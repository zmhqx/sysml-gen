# -*- coding: utf-8 -*-
"""合并文档正文扩充段落（配合五号、单倍行距，目标总篇幅约 40 页）"""
from __future__ import annotations

from docx import Document

from docx_style import add_heading, add_para, add_table


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        add_para(doc, it, indent_cm=0.75)


def _long_para(doc: Document, text: str) -> None:
    add_para(doc, text)


def expand_part1_user_manual(doc: Document) -> None:
    _narrative_block(
        doc,
        "3.0　使用场景说明",
        [
            "本系统主要面向使用 SysML 进行系统建模的毕业设计学生与指导教师，典型场景包括："
            "课程设计文档编制、开题报告中的模型说明章节生成、答辩材料中的设计说明文档导出等。",
            "用户可在本地搭建完整环境后离线使用，亦可在实验室统一服务器部署供多人访问（需自行配置网络与 HTTPS）。",
            "与手工编写 Word 相比，本系统可减少重复拷贝模型属性、统一文档版式，并降低模型变更后的文档同步成本。",
        ],
    )
    add_heading(doc, "3.8　接口与集成说明", 2)
    add_para(
        doc,
        "系统对外主要提供 REST API（基础路径 /api/v1），认证方式为 Bearer JWT。"
        "常用接口分组如下表；完整定义见 Swagger 文档。",
    )
    add_table(
        doc,
        ["模块", "典型接口", "说明"],
        [
            ["认证", "POST /auth/login", "获取 access_token"],
            ["项目", "GET/POST /projects", "项目列表与创建"],
            ["模型", "POST /models/upload", "上传并触发解析"],
            ["模板", "GET/PUT /templates/{id}", "模板查询与更新"],
            ["文档", "POST /documents/generate", "提交生成任务"],
            ["系统", "GET /admin/logs", "操作日志（admin）"],
        ],
        col_widths_cm=[2.5, 4.5, 7.5],
        font_size=10.5,
        align_center=False,
    )

    add_heading(doc, "4.4　用户角色与权限", 2)
    add_table(
        doc,
        ["角色", "项目管理", "模型/模板", "文档生成", "系统管理"],
        [
            ["admin", "全部", "全部", "全部", "全部"],
            ["manager", "所属项目", "所属项目", "所属项目", "只读日志"],
            ["member", "参与项目", "上传/查看", "生成/下载", "无"],
        ],
        col_widths_cm=[2.0, 3.0, 3.0, 3.0, 3.5],
        font_size=10.5,
    )
    add_para(
        doc,
        "权限在服务端中间件统一校验；前端根据角色隐藏无权限菜单。"
        "越权访问将返回 HTTP 403，用户应联系项目 owner 或系统管理员。",
    )

    add_heading(doc, "5.9　模型文件格式说明", 2)
    add_bullets(
        doc,
        [
            "a) 支持 .xmi、.xml、.json 及包含上述文件的 .zip 压缩包；",
            "b) 单文件建议不超过 50MB，超大文件解析时间延长；",
            "c) 解析成功后可在模型详情查看元素列表与 parse 日志；",
            "d) demo_model.json 为验收演示推荐样例。",
        ],
    )

    add_heading(doc, "5.10　文档模板变量说明", 2)
    add_para(
        doc,
        "模板正文采用 HTML + Jinja2 语法，常用变量包括 project、model、elements、"
        "generated_at 等；生成时由后端渲染引擎注入。示例如下表。",
    )
    add_table(
        doc,
        ["变量", "含义", "示例"],
        [
            ["{{ project.name }}", "项目名称", "毕业设计示例"],
            ["{{ model.filename }}", "模型文件名", "demo_model.json"],
            ["{{ elements|length }}", "元素个数", "128"],
            ["{{ generated_at }}", "生成时间", "2026-05-20 14:30"],
        ],
        col_widths_cm=[4.0, 3.5, 6.0],
        font_size=10.5,
        align_center=False,
    )

    add_heading(doc, "5.11　常见问题与处理", 2)
    add_table(
        doc,
        ["现象", "可能原因", "处理步骤"],
        [
            ["无法登录", "口令错误或后端未启动", "检查 uvicorn 与账号状态"],
            ["上传失败", "超 50MB 或格式不支持", "压缩或转换格式后重传"],
            ["解析失败", "模型文件损坏或编码异常", "查看 parse 日志并更换样例"],
            ["生成一直进行中", "PDF 服务异常", "查后端日志，先下载 docx"],
            ["403 无权限", "角色不足", "联系 admin 调整项目成员"],
        ],
        col_widths_cm=[3.0, 4.0, 7.5],
        font_size=10.5,
        align_center=False,
    )
    add_heading(doc, "5.12　界面导航说明", 2)
    add_para(
        doc,
        "登录后左侧导航依次为：首页仪表盘、项目管理、模型管理、模板管理、文档中心、系统管理（仅 admin）。"
        "顶栏显示当前用户、项目切换（如有）与退出；列表页统一提供搜索框、刷新与分页组件。",
    )
    add_para(
        doc,
        "文档中心支持按项目筛选生成记录，点击「生成」打开向导：选择项目 → 选择已解析模型 → 选择模板 → 确认格式 → 提交。"
        "生成完成后可在列表点击「预览」在线查看或「下载」保存到本地。",
    )

    add_heading(doc, "4.5　安装部署详细步骤", 2)
    steps = [
        "1) 安装 MySQL 8.0+，记住 root 口令，创建数据库 sysmldocgen；",
        "2) 在项目 database 目录以管理员执行 init.sql；",
        "3) 复制 backend/.env.example 为 .env，填写 DATABASE_URL 与 SECRET_KEY；",
        "4) 创建 backend/storage 目录并赋予写权限；",
        "5) Python 虚拟环境：python -m venv venv && venv\\Scripts\\activate；",
        "6) pip install -r backend/requirements.txt；",
        "7) 可选：python scripts/seed_templates.py 导入示例模板；",
        "8) 启动后端：cd backend && uvicorn app.main:app --host 0.0.0.0 --port 8000；",
        "9) 另开终端：cd frontend && npm install && npm run dev；",
        "10) 浏览器访问 http://localhost:5173，使用 admin/admin123 登录并修改密码；",
        "11) 上传 demo_model.json 验证解析；",
        "12) 选择模板执行一次文档生成并下载 docx 验证。",
    ]
    add_bullets(doc, steps)


def expand_part2_test_plan(doc: Document) -> None:
    _narrative_block(
        doc,
        "3.5　测试策略说明",
        [
            "测试策略采用“文档审查先行、冒烟验证路径、全面覆盖需求、缺陷驱动回归”的顺序。"
            "文档审查保证需求与设计可测；冒烟测试在 2 日内验证主路径可运行；全面测试按模块并行推进；"
            "每修复一批缺陷即安排回归，避免问题积压。性能与安全测试穿插在全面测试第二阶段执行。",
            "自动化脚本 demo_test.py 用于每日构建后的快速验证，不作为唯一验收依据，人工用例仍为正式记录来源。",
        ],
    )
    add_heading(doc, "4.4　测试通过准则", 2)
    add_table(
        doc,
        ["测试类型", "通过准则"],
        [
            ["功能测试", "关键用例 100% 通过，其余≥95%"],
            ["接口测试", "状态码与响应体符合 SRS"],
            ["性能测试", "上传≤60s，单次生成≤15s"],
            ["安全测试", "无越权与明文口令泄露"],
            ["文档审查", "审查表四项均勾选满足"],
        ],
        col_widths_cm=[3.0, 11.0],
        font_size=10.5,
        align_center=False,
    )

    add_heading(doc, "4.5　测试风险与应对", 2)
    add_table(
        doc,
        ["风险", "影响", "应对措施"],
        [
            ["环境不一致", "中", "固定 TEST-PC-01 配置清单"],
            ["测试数据不足", "中", "维护 demo_model 标准集"],
            ["进度延误", "低", "优先冒烟+核心模块用例"],
            ["缺陷修复不及时", "中", "每日站会跟踪缺陷状态"],
        ],
        col_widths_cm=[3.5, 1.5, 9.5],
        font_size=10.5,
        align_center=False,
    )

    add_heading(doc, "6.1　人员职责明细", 2)
    add_table(
        doc,
        ["姓名", "角色", "主要职责"],
        [
            ["吴一昊", "测试负责人", "计划/报告编制、评审组织"],
            ["申博文", "测试设计", "用例设计、功能与接口执行"],
            ["林文诚", "环境/性能", "环境搭建、性能与安全测试"],
            ["徐乐", "测试执行", "用例执行、缺陷提交与验证"],
            ["张德平", "指导教师", "计划与报告审核批准"],
        ],
        col_widths_cm=[2.0, 2.5, 9.0],
        font_size=10.5,
        align_center=False,
    )

    add_heading(doc, "10　测试交付物", 1)
    add_bullets(
        doc,
        [
            "a) 《软件测试计划》V1.0（本文档）；",
            "b) 《软件测试说明》及 86 条用例总表；",
            "c) 《软件配置项测试报告》V1.0；",
            "d) 测试记录 Excel、缺陷跟踪表及审查表附录 A～C；",
            "e) demo_test.py 执行日志与测试环境截图。",
        ],
    )
    add_heading(doc, "11　测试资源", 1)
    add_para(
        doc,
        "人力资源：测试组 4 人，指导教师 1 人；设备资源：PC 1 台、局域网环境；"
        "软件资源：被测系统一套、MySQL 实例、浏览器与接口测试工具。"
        "上述资源在测试周期内固定使用，避免中途更换导致结果不可比。",
    )

    add_heading(doc, "12　需求可追溯矩阵（摘要）", 1)
    add_table(
        doc,
        ["需求标识", "需求简述", "测试项", "用例前缀"],
        [
            ["SRS-UM-01", "用户登录", "TI-UM", "GN-UM"],
            ["SRS-UM-02", "JWT 鉴权", "TI-UM", "GN-UM/JK"],
            ["SRS-PM-01", "项目 CRUD", "TI-PM", "GN-PM"],
            ["SRS-PM-02", "成员权限", "TI-PM", "GN-PM/AQX"],
            ["SRS-MM-01", "模型上传", "TI-MM", "GN-MM"],
            ["SRS-MM-02", "解析性能", "TI-MM", "XN"],
            ["SRS-TM-01", "模板维护", "TI-TM", "GN-TM"],
            ["SRS-DG-01", "文档生成", "TI-DG", "GN-DG"],
            ["SRS-DG-02", "PDF 导出", "TI-DG", "GN-DG/XN"],
            ["SRS-AD-01", "用户管理", "TI-AD", "GN-AD"],
            ["SRS-AD-02", "操作日志", "TI-AD", "GN-AD"],
            ["SRS-NF-01", "界面易用", "TI-AD", "JM"],
        ],
        col_widths_cm=[2.5, 5.0, 2.0, 4.0],
        font_size=10.5,
        align_center=False,
    )


def expand_part3_test_spec(doc: Document) -> None:
    from docx_style import add_page_break
    from gen_test_cases_md import CASES

    add_heading(doc, "4.5　各模块测试规程摘要", 2)
    modules = [
        (
            "用户认证（TI-UM）",
            "验证登录、登出、令牌刷新、密码修改及未授权拦截；"
            "覆盖 admin/manager/member 三种角色登录场景。",
        ),
        (
            "项目管理（TI-PM）",
            "验证项目 CRUD、成员邀请、权限边界及列表筛选；"
            "重点检查非 owner 用户不能删除他人项目。",
        ),
        (
            "模型管理（TI-MM）",
            "验证 xmi/xml/json/zip 上传、解析成功/失败提示、大文件计时；"
            "检查解析结果元素数量与详情页展示。",
        ),
        (
            "模板与文档（TI-TM/DG）",
            "验证模板保存、变量渲染、Word 生成与 PDF 导出；"
            "检查生成文件可打开且中文显示正常。",
        ),
        (
            "系统管理（TI-AD）",
            "验证用户管理、角色分配、日志查询与 admin 专属接口保护。",
        ),
    ]
    for title, desc in modules:
        add_heading(doc, title, 3)
        add_para(doc, desc)
        prefix_map = {
            "用户认证": "GN-UM",
            "项目管理": "GN-PM",
            "模型管理": "GN-MM",
            "模板与文档": "GN-TM",
            "系统管理": "GN-AD",
        }
        key = title.split("（")[0]
        pf = prefix_map.get(key, "GN")
        ids = [c["id"] for c in CASES if c["id"].startswith(pf)][:10]
        add_para(doc, "相关用例（节选）：" + "、".join(ids) + "。")

    add_heading(doc, "4.6　回归测试说明", 2)
    add_para(
        doc,
        "每轮缺陷修复后，执行缺陷关联用例及冒烟集共 12 条；"
        "全部通过后方可关闭缺陷并进入下一轮全面测试。",
    )
    add_heading(doc, "4.7　测试记录要求", 2)
    add_para(
        doc,
        "每条用例执行时应记录：执行日期、执行人、实际输出、测试结论（通过/失败/阻塞）、"
        "失败时附截图或日志片段。阻塞须注明原因（环境、数据、代码未就绪等）并升级至测试负责人。",
    )
    add_heading(doc, "4.8　文档审查执行说明", 2)
    add_bullets(
        doc,
        [
            "a) 依据附录 A 审查《软件需求规格说明》，确认需求可测、无矛盾；",
            "b) 依据附录 B 审查《软件概要设计说明》，确认设计与需求可追溯；",
            "c) 依据附录 C 审查《软件用户手册》，确认安装与操作步骤可复现；",
            "d) 审查结论填入审查表，签字后纳入测试报告附件。",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "附录 B　用例执行要点索引", 1)
    add_para(
        doc,
        "以下按用例标识列出执行要点摘要，完整步骤与预期输出以附录 A-1 总表及代表详表为准。",
    )
    for c in CASES:
        add_para(
            doc,
            f"{c['id']}（{c['type']}）{c['name']}：前置—{c['pre'][:36]}；步骤—{c['steps'][:48]}；预期—{c['exp'][:36]}。",
            indent_cm=0.75,
        )

    add_heading(doc, "4.9　轮次测试安排", 2)
    add_table(
        doc,
        ["轮次", "时间", "用例数", "重点"],
        [
            ["冒烟", "04-17～18", "15", "登录、项目、上传、生成主路径"],
            ["全面-I", "04-19～04-30", "35", "各模块功能与接口"],
            ["全面-II", "05-01～05-10", "36", "性能、安全、边界、界面"],
            ["回归", "05-11～18", "12", "缺陷关联与冒烟复核"],
        ],
        col_widths_cm=[2.0, 3.5, 2.0, 6.0],
        font_size=10.5,
        align_center=False,
    )


def _narrative_block(doc: Document, title: str, paragraphs: list[str]) -> None:
    add_heading(doc, title, 2)
    for t in paragraphs:
        add_para(doc, t)


def expand_part4_test_report(doc: Document) -> None:
    add_heading(doc, "3.5　测试工具与数据", 2)
    add_table(
        doc,
        ["工具/数据", "版本或说明", "用途"],
        [
            ["Chrome", "120+", "界面与 E2E 操作"],
            ["Swagger UI", "内置", "接口调试"],
            ["demo_test.py", "项目脚本", "冒烟自动化"],
            ["demo_model.json", "标准样例", "解析与生成测试"],
            ["mysqldump", "8.4", "测试数据备份"],
        ],
        col_widths_cm=[3.0, 3.5, 7.0],
        font_size=10.5,
        align_center=False,
    )

    add_heading(doc, "5.7　缺陷清单摘要", 2)
    add_table(
        doc,
        ["编号", "模块", "摘要", "等级", "状态"],
        [
            ["BUG-001", "认证", "令牌过期无友好提示", "S3", "已关闭"],
            ["BUG-002", "项目", "删除后列表未刷新", "S3", "已关闭"],
            ["BUG-003", "模型", "大文件解析超时", "S2", "已关闭"],
            ["BUG-004", "文档", "PDF 中文乱码", "S2", "已关闭"],
            ["BUG-005", "界面", "侧边栏折叠异常", "S4", "已关闭"],
            ["BUG-012", "界面", "顶栏搜索无后端", "S4", "遗留"],
        ],
        col_widths_cm=[1.8, 2.0, 5.5, 1.2, 1.5],
        font_size=10.5,
        align_center=False,
    )

    add_heading(doc, "6.3　质量评价", 2)
    add_para(
        doc,
        "从功能性、可靠性、易用性、效率及可维护性五个方面评价："
        "核心功能完整可用；关键缺陷已修复；界面风格统一；"
        "在实验室环境下性能达标；代码与文档结构清晰，便于后续扩展。",
    )
    add_heading(doc, "6.4　测试完成度", 2)
    add_table(
        doc,
        ["项目", "计划", "实际", "说明"],
        [
            ["用例设计", "86 条", "86 条", "与 SRS 追溯完整"],
            ["用例执行", "86 条", "86 条", "执行率 100%"],
            ["缺陷关闭", "—", "11/12", "遗留 1 项 S4"],
            ["文档审查", "3 份", "3 份", "附录 A～C 合格"],
        ],
        col_widths_cm=[3.0, 2.5, 2.5, 6.5],
        font_size=10.5,
        align_center=False,
    )
    add_heading(doc, "8.1　术语", 2)
    add_table(
        doc,
        ["术语", "说明"],
        [
            ["冒烟测试", "对核心功能快速验证"],
            ["回归测试", "缺陷修复后的再测试"],
            ["配置项测试", "CSCI 级测试"],
            ["缺陷密度", "缺陷数/千行代码或功能点"],
        ],
        col_widths_cm=[3.0, 11.5],
        font_size=10.5,
    )

    add_heading(doc, "8.2　测试过程纪要", 2)
    add_para(
        doc,
        "第一轮冒烟测试于 4 月 17 日开始，当日完成环境搭建与 15 条用例，暴露登录态刷新与模型列表分页问题；"
        "4 月 19 日起进入全面测试，按模块分工：申博文负责 PM/TM，林文诚负责 MM/XN，徐乐负责 UM/AD/JK；"
        "5 月 11 日启动回归，重点验证 BUG-003、BUG-004 修复效果，至 5 月 18 日全部通过。",
    )
    add_para(
        doc,
        "测试期间每日 17:00 召开 15 分钟站会同步进度，缺陷统一登记至跟踪表并由吴一昊每日汇总发送指导教师。"
        "5 月 19 日起编写本报告，5 月 24 日经张德平老师审核批准。",
    )
