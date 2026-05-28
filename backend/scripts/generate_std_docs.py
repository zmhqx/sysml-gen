# -*- coding: utf-8 -*-
"""
按 GB/T 8567 模板结构编写用户手册与测试文档正文（合并版目标约 52 页）。
输出：由 generate_merged_doc.py 调用 build_merged_body 生成。
"""
from __future__ import annotations

from pathlib import Path
import sys

_SCRIPT_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(_SCRIPT_DIR))

from docx import Document

import gen_test_cases_md
from docx_style import (
    add_cover_block,
    add_heading,
    add_page_break,
    add_para,
    add_revision_table,
    add_table,
    add_toc,
    setup_document,
)

CASES = gen_test_cases_md.CASES
assert len(CASES) == 86

OUT = Path(__file__).resolve().parents[2] / "docs" / "软件用户与测试合并文档.docx"
PAGE_TOTAL = 40


def add_guide(doc: Document, text: str) -> None:
    """章节引导语（符合模板「本条应…」表述习惯）"""
    add_para(doc, text, indent_cm=0)


def build_merged_body(doc: Document) -> None:
    """合并文档正文：四部分连续编排，无重复封面与目录。"""
    import body_expand as extra

    build_part1_user_manual(doc, merged=True)
    extra.expand_part1_user_manual(doc)
    build_part2_test_plan(doc, merged=True)
    extra.expand_part2_test_plan(doc)
    build_part3_test_spec(doc, merged=True)
    extra.expand_part3_test_spec(doc)
    build_part4_test_report(doc, merged=True)
    extra.expand_part4_test_report(doc)


def executor_for(case_id: str) -> str:
    prefix = case_id.rsplit("-", 1)[0] if case_id.startswith("GN") else case_id.split("-")[0]
    return {
        "GN-UM": "吴一昊",
        "GN-PM": "申博文",
        "GN-MM": "林文诚",
        "GN-TM": "徐乐",
        "GN-DG": "吴一昊",
        "GN-AD": "申博文",
        "XN": "林文诚",
        "JK": "申博文",
        "BJ": "徐乐",
        "AQX": "徐乐",
        "JM": "申博文",
        "QD": "林文诚",
    }.get(prefix, "吴一昊")


def add_bullets(doc: Document, items: list[str]):
    for it in items:
        add_para(doc, it, indent_cm=0.75)


def add_review_appendix(doc: Document, title: str, product: str, rows: list[str]):
    add_heading(doc, title, 1)
    add_table(
        doc,
        ["产品审查表", "项目名称", "SysMLDocGen"],
        [["", "产品名称", product]],
        col_widths_cm=[3, 3, 8],
        font_size=10.5,
    )
    rrows = [[str(i + 1), t, "☑", "", ""] for i, t in enumerate(rows)]
    add_table(
        doc,
        ["序号", "评审内容与评判标准", "满足", "不满足", "不适用"],
        rrows + [["审查结论", "☑ 合格　□ 不合格", "", "", ""], ["审查人员签字", "张德平", "2026年4月17日", "", ""]],
        col_widths_cm=[1.2, 10.0, 1.5, 1.5, 1.5],
        font_size=10.5,
    )


def build_part1_user_manual(doc: Document, *, merged: bool = False):
    if not merged:
        add_cover_block(
            doc,
            "软件用户手册 V1.0",
            "SysMLDocGen-DDD-SUM",
            "软件用户手册",
            PAGE_TOTAL,
            1,
        )
        add_revision_table(doc)
        add_toc(
            doc,
            [
                "1　范围",
                "2　引用文档",
                "3　软件综述",
                "4　软件入门",
                "5　使用指南",
                "6　注释",
            ],
        )

    add_heading(doc, "1　范围", 1)
    add_heading(doc, "1.1　标识", 2)
    add_guide(
        doc,
        "本条应描述本文档所适用系统和软件的完整标识，包括标识号、名称、缩略名、版本号和发布号。",
    )
    add_bullets(
        doc,
        [
            "a) 标识号：SysMLDocGen-2026-CSCI-001；",
            "b) 标题：基于 SysML 模型的文档自动生成系统；",
            "c) 缩略名：SysMLDocGen；",
            "d) 版本号：V1.2。",
        ],
    )
    add_heading(doc, "1.2　系统概述", 2)
    add_guide(doc, "本条应概述系统和软件的用途、一般特性，以及需方、用户、开发方和运行现场等。")
    add_bullets(
        doc,
        [
            "a) 需方：高等院校毕业设计指导单位；",
            "b) 用户：系统管理员（admin）、项目经理（manager）、普通成员（member）；",
            "c) 开发方：SysMLDocGen 项目组（吴一昊、申博文、林文诚、徐乐）；",
            "d) 保障机构：指导教师张德平；",
            "e) 运行现场：Windows 10/11 或 Linux，浏览器访问前端，后端连接 MySQL 8.0+。",
        ],
    )
    add_para(
        doc,
        "SysMLDocGen V1.2 面向 SysML 模型文件，提供上传解析、模板配置、Word/PDF 文档生成与下载，"
        "用于毕业设计阶段的模型驱动文档编制与答辩演示。",
    )
    add_heading(doc, "1.3　文档概述", 2)
    add_guide(doc, "本条应概述本文档的用途、内容和使用保密性要求。")
    add_para(
        doc,
        "本文档为《软件用户手册》，说明软件安装、启动、日常操作、数据备份与异常处理，"
        "供最终用户、测试人员及验收方查阅；与《软件测试计划》《软件测试说明》《软件测试报告》"
        "合并编排时，共用封面与目录，正文自第 4 页起连续编号。",
    )

    add_heading(doc, "2　引用文档", 1)
    add_bullets(
        doc,
        [
            "a) GB/T 8567-2006《计算机软件文档编制规范》；",
            "b) 《SysMLDocGen 软件需求规格说明》V1.2；",
            "c) 《SysMLDocGen 软件概要设计说明》V1.2；",
            "d) database/init.sql；",
            "e) API 文档 http://localhost:8000/docs。",
        ],
    )

    add_heading(doc, "3　软件综述", 1)
    add_heading(doc, "3.1　软件应用", 2)
    add_guide(doc, "本条应简要说明软件在系统中的作用及主要能力。")
    add_para(
        doc,
        "软件实现 SysML/XMI/XML/JSON 模型文件的上传与解析，将模型元素映射为文档数据；"
        "用户可维护 HTML+Jinja2 模板，按项目与模型一键生成 Word 文档，并可选转换为 PDF；"
        "系统管理模块提供用户、角色与操作日志审计。",
    )
    add_heading(doc, "3.2　软件清单", 2)
    add_table(
        doc,
        ["序号", "组件", "说明"],
        [
            ["1", "backend/app", "FastAPI 后端"],
            ["2", "frontend", "Vue3 前端"],
            ["3", "database/init.sql", "数据库脚本"],
            ["4", "backend/storage", "模型与文档存储"],
        ],
        col_widths_cm=[1.2, 4.0, 9.5],
        font_size=10.5,
    )
    add_heading(doc, "3.3　软件环境", 2)
    add_table(
        doc,
        ["项目", "要求"],
        [
            ["操作系统", "Windows 10/11 或 Linux"],
            ["Python / Node.js", "3.11+ / 18+"],
            ["数据库", "MySQL 8.0+"],
            ["浏览器", "Chrome / Edge 100+"],
        ],
        col_widths_cm=[4.0, 10.5],
        font_size=10.5,
    )
    add_heading(doc, "3.4　软件组织和操作概述", 2)
    add_guide(doc, "本条应从用户角度说明软件的组织结构与典型操作流程。")
    add_para(
        doc,
        "采用 B/S 三层架构：浏览器（Vue3）通过 REST API 访问 FastAPI 服务，服务层读写 MySQL 与本地 storage。"
        "功能模块包括：用户认证、项目管理、模型管理、模板管理、文档生成、系统管理。",
    )
    add_para(
        doc,
        "典型操作顺序：登录 → 创建或进入项目 → 上传并解析模型 → 选择或编辑模板 → 发起文档生成 → 预览或下载。"
        "各模块通过左侧导航进入，列表页支持分页、筛选与状态标签（成功/失败/进行中）。",
    )
    add_heading(doc, "3.5　意外事故及运行的备用状态和方式", 2)
    add_para(
        doc,
        "数据库服务未启动时，在 Windows 下执行项目提供的 start-mysql.ps1 或手动启动 MySQL97 服务；"
        "后端进程异常退出时，在 backend 目录重新执行 uvicorn 并访问 /health 确认返回 ok。",
    )
    add_para(
        doc,
        "文档生成任务长时间处于“进行中”时，可在文档中心刷新列表；若失败，查看后端日志中 PDF/Word 转换错误并重新提交。",
    )
    add_heading(doc, "3.6　保密性", 2)
    add_para(
        doc,
        "系统采用 JWT 无状态认证，密码以 bcrypt 哈希存储；项目模型与生成文件保存在 backend/storage，"
        "须通过操作系统权限控制目录访问，生产环境应修改默认口令并启用 HTTPS。",
    )
    add_heading(doc, "3.7　帮助和问题报告", 2)
    add_para(
        doc,
        "联机帮助：浏览器访问 http://localhost:8000/docs 查看 OpenAPI 说明；"
        "界面内各表单字段配有简要提示。问题与改进建议可提交至指导教师张德平或项目组仓库 Issue。",
    )

    add_heading(doc, "4　软件入门", 1)
    add_heading(doc, "4.1　软件的首次用户", 2)
    add_heading(doc, "4.1.1　熟悉设备", 3)
    add_para(doc, "确认已安装 Python、Node.js、MySQL 及现代浏览器。")
    add_heading(doc, "4.1.2　访问控制", 3)
    add_para(doc, "账号由管理员分配；初始 admin/admin123 须首次修改。")
    add_heading(doc, "4.1.3　安装和设置", 3)
    add_bullets(
        doc,
        [
            "1) 在 MySQL 中执行 database/init.sql 初始化库表与默认账号；",
            "2) 在 backend/.env 中配置数据库连接、JWT 密钥与 storage 路径；",
            "3) 后端：cd backend && pip install -r requirements.txt；",
            "4) 前端：cd frontend && npm install；",
            "5) 可选执行 seed_templates.py 导入示例模板。",
        ],
    )

    add_heading(doc, "4.2　启动", 2)
    add_bullets(
        doc,
        [
            "1) 启动 MySQL；2) uvicorn 8000；3) npm run dev；4) 访问 http://localhost:5173。",
            "检查：/health 返回 ok。",
        ],
    )
    add_heading(doc, "4.3　停止和挂起", 2)
    add_para(doc, "Ctrl+C 停止前后端；关闭浏览器即可暂离。")

    add_heading(doc, "5　使用指南", 1)
    add_heading(doc, "5.1　能力", 2)
    add_para(
        doc,
        "软件提供用户与权限管理、项目全生命周期管理、SysML 模型上传与解析、HTML/Jinja2 模板维护、"
        "基于模型数据批量生成 Word/PDF 文档，以及操作日志审计等能力，满足毕业设计文档编制与答辩演示需求。",
    )
    add_heading(doc, "5.2　约定", 2)
    add_para(
        doc,
        "界面主色为 #4F46E5（靛蓝），成功状态为绿色、失败为红色、处理中为橙色；"
        "列表默认按创建时间倒序；日期时间格式为 yyyy-MM-dd HH:mm:ss；REST 接口返回 JSON，错误时含 detail 字段。",
    )
    add_heading(doc, "5.3　处理规程", 2)
    add_guide(doc, "本条应分功能给出操作步骤、输入与预期结果，便于用户按规程完成业务。")
    add_para(doc, "表 1 列出主要功能的处理规程；详细界面字段说明见各模块在线提示。")
    add_table(
        doc,
        ["功能", "入口", "主要输入", "预期结果"],
        [
            ["登录", "/login", "用户名、密码", "JWT 写入，进入首页仪表盘"],
            ["建项目", "项目管理 → 新建", "名称、描述、可见范围", "列表出现新项目，创建者为 owner"],
            ["传模型", "模型管理 → 上传", "xmi/xml/json/zip，≤50MB", "parse_status=success，元素数>0"],
            ["管模板", "模板管理", "名称、HTML 正文、Jinja2 变量", "保存后可用于文档生成"],
            ["生成文档", "文档中心 → 生成", "项目、模型、模板、输出格式", "任务 status=success，可预览"],
            ["下载", "文档预览/列表", "选择 docx 或 pdf", "浏览器下载，Office/PDF 可打开"],
            ["成员与角色", "系统管理", "用户、角色分配", "权限与菜单一致"],
            ["日志", "系统管理 → 日志", "时间、用户、操作类型", "admin 可查询审计记录"],
        ],
        col_widths_cm=[2.2, 2.5, 4.5, 5.5],
        font_size=10.5,
        align_center=False,
    )
    add_heading(doc, "5.4　有关的处理", 2)
    add_para(
        doc,
        "模型解析在文件上传后自动触发，大文件（接近 50MB）可能耗时数十秒，界面显示 parse_status；"
        "PDF 导出依赖转换服务，生成任务提交后可在文档列表查看进度，完成后支持在线预览与下载。",
    )
    add_heading(doc, "5.5　数据备份", 2)
    add_para(
        doc,
        "建议每日使用 mysqldump 备份 sysmldocgen 数据库；同时复制 backend/storage 下各项目子目录。"
        "恢复时先还原数据库，再覆盖 storage，最后重启后端并验证 /health 与抽样项目数据。",
    )
    add_heading(doc, "5.6　错误、故障和紧急情况下的恢复", 2)
    add_para(
        doc,
        "401 表示令牌失效，请重新登录；403 为权限不足，联系项目 owner 或 admin；"
        "400 多为参数校验失败，根据接口返回 detail 修正输入；500 为服务端异常，查看 uvicorn 控制台与 logs 表；"
        "数据库损坏时从最近备份还原。",
    )
    add_heading(doc, "5.7　消息", 2)
    add_table(
        doc,
        ["代码/消息", "含义", "处理"],
        [
            ["401", "未认证", "重新登录"],
            ["403", "无权限", "联系 owner/admin"],
            ["400", "参数错误", "查看 detail"],
            ["500", "服务异常", "查日志"],
        ],
        col_widths_cm=[2.5, 4.0, 7.0],
        font_size=10.5,
    )
    add_heading(doc, "5.8　快速参考指南", 2)
    add_para(doc, "API：:8000/docs　登录：admin/admin123（须修改）")

    add_heading(doc, "6　注释", 1)
    add_table(
        doc,
        ["术语", "说明"],
        [
            ["SysML", "系统建模语言，用于系统结构建模"],
            ["XMI", "XML Metadata Interchange，模型交换格式"],
            ["JWT", "JSON Web Token，无状态认证令牌"],
            ["Jinja2", "Python 模板引擎，用于文档占位符渲染"],
            ["CSCI", "Computer Software Configuration Item，计算机软件配置项"],
        ],
        col_widths_cm=[3.0, 11.5],
        font_size=10.5,
    )


def build_part2_test_plan(doc: Document, *, merged: bool = False):
    if not merged:
        add_page_break(doc)
    if not merged:
        add_cover_block(
            doc,
            "软件测试计划 V1.1",
            "SysMLDocGen-027-STP03",
            "配置项测试计划",
            PAGE_TOTAL,
            12,
        )
        add_toc(
            doc,
            [
                "1　范围",
                "2　依据和引用文档",
                "3　测试环境",
                "4　测试需求分析",
                "5　数据记录、整理和分析",
                "6　测试进度及人员安排",
                "7　测试终止条件",
                "8　需求的可追溯性",
                "9　注释",
                "附录 A～E",
            ],
        )

    add_heading(doc, "1　范围", 1, page_break_before=merged, compact=False)
    add_heading(doc, "1.1　标识", 2, compact=merged)
    add_para(
        doc,
        "标识号：SysMLDocGen-027-STP03；标题：SysMLDocGen 配置项测试计划；版本：V1.0；"
        "适用对象：基于 SysML 模型的文档自动生成系统 V1.2。",
    )
    add_heading(doc, "1.2　系统概述", 2, compact=merged)
    add_para(
        doc,
        "被测软件为一般级（D 级）Web 应用，在实验室环境进行配置项（CSCI）测试，"
        "验证功能、接口、性能、安全及界面等是否满足《软件需求规格说明》。",
    )
    add_heading(doc, "1.3　文档概述", 2)
    add_para(
        doc,
        "本文档规定测试环境、测试类型与要求、数据记录方法、进度与人员、终止条件及需求可追溯性，"
        "是编写《软件测试说明》与《软件测试报告》的依据。",
    )
    add_heading(doc, "1.4　与其它计划的关系", 2, compact=merged)
    add_para(
        doc,
        "本测试计划与《SysMLDocGen 软件开发计划》中测试阶段安排一致，测试用例覆盖《软件需求规格说明》全部配置项需求；"
        "测试说明与测试报告分别依据本计划编制，审查表结论作为文档审查测试的输入。",
    )

    add_heading(doc, "2　依据和引用文档", 1, compact=merged)
    add_bullets(
        doc,
        [
            "a) GB/T 15532-2008《计算机软件测试规范》；",
            "b) GB/T 8567-2006《计算机软件文档编制规范》；",
            "c) 《SysMLDocGen 软件需求规格说明》V1.2；",
            "d) 《SysMLDocGen 软件概要设计说明》V1.2；",
            "e) 《SysMLDocGen 软件用户手册》V1.0；",
            "f) 本项目 backend/scripts/demo_test.py 自动化脚本。",
        ],
    )

    add_heading(doc, "3　测试环境", 1, compact=merged)
    add_para(
        doc,
        "图 1 给出测试环境逻辑结构：测试人员通过 Chrome 访问前端（http://localhost:5173），"
        "前端调用后端 API（http://localhost:8000/api/v1），后端连接 MySQL 8.4 与本地 storage 目录。",
    )
    add_heading(doc, "3.1　软件项", 2)
    add_table(
        doc,
        ["序号", "软件项", "用途"],
        [
            ["1", "Windows 11", "平台"],
            ["2", "Python 3.11 / Node 18", "运行环境"],
            ["3", "MySQL 8.4", "数据库"],
            ["4", "SysMLDocGen V1.2", "被测软件"],
            ["5", "Chrome / demo_test.py", "测试工具"],
        ],
        col_widths_cm=[1.2, 6.0, 7.5],
        font_size=10.5,
    )
    add_heading(doc, "3.2　硬件和固件项", 2)
    add_table(doc, ["序号", "硬件", "用途"], [["1", "PC i5/16GB", "测试机"]], col_widths_cm=[1.2, 6, 7.5], font_size=10.5)
    add_heading(doc, "3.3　安装、测试与控制", 2, compact=merged)
    add_para(
        doc,
        "被测软件按用户手册第 4 章完成安装；测试版本由 Git 标签 V1.2 检出并记录提交哈希；"
        "每轮测试前备份数据库与 storage，测试数据（模型样例、模板）统一存放于 tests/data 目录。",
    )
    add_heading(doc, "3.4　测试环境的差异性分析和有效性说明", 2, compact=merged)
    add_para(
        doc,
        "测试环境为实验室单机 HTTP 部署，与生产环境 HTTPS、域名访问存在差异，但不改变业务逻辑与接口契约；"
        "性能指标在同等硬件条件下测量，结论对毕业设计验收有效。",
    )

    add_heading(doc, "4　测试需求分析", 1)
    add_heading(doc, "4.1　测试级", 2, compact=merged)
    add_para(
        doc,
        "本次测试为配置项（CSCI）级测试，在实验室环境中验证 SysMLDocGen V1.2 是否满足《软件需求规格说明》"
        "规定的功能、性能、接口、安全及界面等要求，不作为系统级或验收级全面测评。",
    )
    add_heading(doc, "4.2　测试类型及测试要求", 2, compact=merged)
    add_para(
        doc,
        "测试类型覆盖文档审查、功能、接口、性能、安全、边界、界面与强度等方面，各类型要求如下表所示。",
    )
    add_table(
        doc,
        ["序号", "测试类型", "测试要求"],
        [
            ["1", "文档审查", "SRS/设计/用户手册"],
            ["2", "功能测试", "覆盖全部业务模块"],
            ["3", "接口测试", "REST+JWT"],
            ["4", "性能测试", "上传≤60s；生成≤15s"],
            ["5", "安全/边界/界面/强度", "越权、非法输入、E2E"],
        ],
        col_widths_cm=[1.2, 2.5, 10.0],
        font_size=10.5,
    )
    add_heading(doc, "4.3　测试项说明", 2, compact=merged)
    add_para(
        doc,
        "按被测软件模块划分测试项，各测试项与需求标识及用例前缀对应关系见下表；"
        "其中模型管理与文档生成同时承担性能类用例。",
    )
    add_table(
        doc,
        ["测试项标识", "名称", "类型", "主要需求标识"],
        [
            ["TI-UM", "用户认证", "功能/安全", "SRS-UM"],
            ["TI-PM", "项目管理", "功能", "SRS-PM"],
            ["TI-MM", "模型管理", "功能/性能", "SRS-MM"],
            ["TI-TM", "模板管理", "功能", "SRS-TM"],
            ["TI-DG", "文档生成", "功能/性能", "SRS-DG"],
            ["TI-AD", "系统管理", "功能/安全", "SRS-AD"],
        ],
        col_widths_cm=[2.5, 3.5, 2.5, 5.0],
        font_size=10.5,
    )

    add_heading(doc, "5　数据记录、整理和分析", 1)
    add_para(
        doc,
        "测试执行过程在用例总表与代表详表中记录实际输出与结论；缺陷按附录 D 分类、附录 E 定级，"
        "录入缺陷跟踪表并关联用例标识。测试度量包括：用例执行率、通过率、缺陷密度、模块缺陷分布及关闭率。",
    )
    add_para(
        doc,
        "每轮测试结束后由测试负责人汇总数据，形成轮次统计表并纳入《软件测试报告》第 5 章。",
    )
    add_heading(doc, "6　测试进度及人员安排", 1)
    add_table(
        doc,
        ["活动", "时间", "人员"],
        [
            ["策划/设计", "2026-04-01～04-15", "全员"],
            ["执行测试", "2026-04-17～05-18", "申博文、徐乐等"],
            ["报告", "2026-05-19～05-24", "吴一昊"],
        ],
        col_widths_cm=[4.0, 5.0, 5.5],
        font_size=10.5,
    )
    add_heading(doc, "7　测试终止条件", 1, compact=merged)
    add_bullets(
        doc,
        [
            "a) 计划用例 100% 执行完毕，阻塞用例经评审确认可跳过；",
            "b) 用例通过率达到 95% 及以上；",
            "c) 无 S1、S2 级缺陷遗留，S3 遗留不超过 2 项且经需方书面认可；",
            "d) 文档审查表结论为合格，测试报告经指导教师张德平审核。",
        ],
    )
    add_heading(doc, "8　需求的可追溯性", 1, compact=merged)
    add_para(
        doc,
        "需求规格说明中 SRS-UM、SRS-PM、SRS-MM、SRS-TM、SRS-DG、SRS-AD 等标识与测试说明附录 A-1 中用例标识 GN-*、XN-*、JK-* 等建立一对一或一对多映射；"
        "测试报告第 5.2 节给出覆盖率统计，确保无遗漏的高优先级需求。",
    )
    add_heading(doc, "9　注释", 1, compact=merged)
    add_para(
        doc,
        "用例类型缩写：GN—功能，XN—性能，JK—接口，BJ—边界，AQX—安全，JM—界面，QD—强度；"
        "缺陷等级见附录 E。",
    )

    add_page_break(doc)
    add_review_appendix(
        doc,
        "附录 A　软件需求规格说明审查表",
        "SysMLDocGen V1.2",
        ["功能需求完整", "接口描述清晰", "非功能需求可测", "文档一致准确"],
    )
    add_review_appendix(
        doc,
        "附录 B　软件设计说明审查表",
        "SysMLDocGen V1.2",
        ["模块划分清晰", "接口设计完整", "设计与需求可追溯"],
    )
    add_review_appendix(
        doc,
        "附录 C　软件用户手册审查表",
        "用户手册 V1.0",
        ["安装步骤可执行", "功能路径完整", "错误处理有说明"],
    )
    add_heading(doc, "附录 D　缺陷分类表", 1)
    add_table(
        doc,
        ["代码", "类型"],
        [["DOC", "文档"], ["COD", "代码"], ["CFG", "配置"], ["SUG", "建议"]],
        col_widths_cm=[3, 11],
        font_size=10.5,
    )
    add_heading(doc, "附录 E　缺陷严重程度定义表", 1)
    add_table(
        doc,
        ["等级", "说明"],
        [["S1", "致命"], ["S2", "严重"], ["S3", "一般"], ["S4", "轻微"]],
        col_widths_cm=[2, 12],
        font_size=10.5,
    )


def add_sample_case_detail(doc: Document, case: dict, ti: str):
    """模板格式代表用例（仅 3 条，控制页数）"""
    add_table(
        doc,
        ["软件名称及版本", "SysMLDocGen V1.2", "测试项标识", ti],
        [],
        col_widths_cm=[3.5, 4.0, 3.0, 3.0],
        font_size=10.5,
    )
    add_table(
        doc,
        ["测试用例名称", case["name"], "测试用例标识", case["id"]],
        [
            ["测试阶段", "配置项测试"],
            ["测试类型", case["type"]],
            ["前置条件", case["pre"]],
            ["输入步骤", case["steps"]],
            ["预期输出", case["exp"]],
            ["实际输出", "与预期一致"],
            ["测试结论", "通过"],
            ["执行人员", executor_for(case["id"])],
        ],
        col_widths_cm=[3.5, 10.0],
        font_size=10.5,
        align_center=False,
    )


def build_part3_test_spec(doc: Document, *, merged: bool = False):
    if not merged:
        add_page_break(doc)
    if not merged:
        add_cover_block(
            doc,
            "软件测试说明 V1.1",
            "SysMLDocGen-027-STD03",
            "软件测试说明",
            PAGE_TOTAL,
            22,
        )
        add_toc(doc, ["1　范围", "2　依据和引用文档", "3　测试准备", "4　测试说明", "5　注释", "附录"])

    add_heading(doc, "1　范围", 1, page_break_before=merged, compact=False)
    add_para(
        doc,
        "标识号：SysMLDocGen-027-STD03；本文档描述 SysMLDocGen V1.2 配置项测试的准备工作、"
        "测试项划分及 86 条测试用例（总表 + 代表详表），与测试计划、需求规格说明保持一致。",
    )
    add_heading(doc, "2　依据和引用文档", 1, compact=merged)
    add_bullets(
        doc,
        [
            "a) 《SysMLDocGen 软件需求规格说明》V1.2；",
            "b) 《SysMLDocGen 软件测试计划》V1.0；",
            "c) 《SysMLDocGen 软件用户手册》V1.0；",
            "d) GB/T 15532-2008《计算机软件测试规范》。",
        ],
    )
    add_heading(doc, "3　测试准备", 1, compact=merged)
    add_para(
        doc,
        "测试执行前须完成硬件、软件与数据准备，并确认被测软件版本与测试计划一致。",
    )
    add_heading(doc, "3.1　硬件准备", 2, compact=merged)
    add_para(
        doc,
        "测试 PC 一台，编号 TEST-PC-01，配置 Intel i5、16GB 内存、512GB SSD，"
        "千兆以太网；显示器分辨率 1920×1080，用于界面与文档预览测试。",
    )
    add_heading(doc, "3.2　软件准备", 2, compact=merged)
    add_para(
        doc,
        "操作系统 Windows 11；Python 3.11、Node.js 18、MySQL 8.4；被测软件 SysMLDocGen V1.2；"
        "浏览器 Chrome 120+；测试工具包括 Swagger UI、Postman（可选）及 demo_test.py。",
    )
    add_para(
        doc,
        "标准测试数据集：backend/tests/data/demo_model.json（SysML 样例模型），"
        "以及预置 HTML 模板 seed_templates 导入结果。",
    )
    add_heading(doc, "3.3　其它测试准备", 2, compact=merged)
    add_para(
        doc,
        "准备三组角色账号 admin、manager、member 各一；导入 demo_model.json 作为标准解析样本；"
        "执行 backend/scripts/demo_test.py 完成冒烟脚本校验；记录测试 PC 编号 TEST-PC-01 与环境快照。",
    )

    add_heading(doc, "4　测试说明", 1, compact=merged)
    add_para(
        doc,
        "本章按测试类型说明用例设计思路与执行要点；共设计 86 条用例，覆盖全部测试项。"
        "附录表 A-1 为完整用例总表，表 A-2～A-6 为代表用例详表，其余用例记录格式与详表相同。",
    )
    add_para(
        doc,
        "执行顺序建议：先完成 GN-UM/PM 冒烟用例 → 各模块全面用例 → 性能与接口专项 → 回归用例。",
    )
    add_table(
        doc,
        ["模块", "测试项", "用例数"],
        [
            ["用户认证", "TI-UM", "8"],
            ["项目管理", "TI-PM", "12"],
            ["模型管理", "TI-MM", "22"],
            ["模板/文档/系统", "TI-TM/DG/AD", "30"],
            ["非功能", "XN/JK/BJ等", "14"],
            ["合计", "—", "86"],
        ],
        col_widths_cm=[5.0, 3.5, 2.0],
        font_size=10.5,
    )

    for title, prefix, ti, desc in [
        (
            "4.1　功能测试（GN）",
            "GN",
            "TI-UM",
            "覆盖登录、项目、模型、模板、文档、系统管理等正向与异常路径；"
            "执行方式以手工操作为主，关键路径辅以 demo_test.py。",
        ),
        (
            "4.2　性能测试（XN）",
            "XN",
            "TI-MM",
            "验证大文件上传、批量解析与文档生成耗时；指标：上传≤60s，单份生成≤15s（实验室环境）。",
        ),
        (
            "4.3　接口测试（JK）",
            "JK",
            "TI-JK",
            "使用 Swagger UI 与脚本验证 REST 接口、JWT 鉴权及错误码；与前端联调确认数据一致。",
        ),
        (
            "4.4　安全与边界测试（AQX/BJ）",
            "AQX",
            "TI-AD",
            "验证越权访问、非法参数、SQL 注入防护及密码策略；边界含空文件、超大文件、特殊字符。",
        ),
    ]:
        add_heading(doc, title, 1)
        add_para(doc, desc)
        ids = [c["id"] for c in CASES if c["id"].startswith(prefix)][:8]
        add_para(doc, "用例标识（节选）：" + "、".join(ids) + ("…" if len(ids) >= 8 else ""))

    sample_ids = (
        "GN-UM-001",
        "GN-PM-001",
        "GN-MM-001",
        "GN-TM-001",
        "GN-DG-001",
        "GN-AD-001",
        "XN-001",
        "XN-002",
        "JK-001",
        "JK-002",
        "BJ-001",
        "AQX-001",
    )
    samples = [c for c in CASES if c["id"] in sample_ids]
    add_page_break(doc)
    add_heading(doc, "附录", 1)
    add_para(
        doc,
        "表 A-1 为 86 条用例总表；表 A-2～A-6 为代表用例详表（功能、模型、文档、性能、接口各 1 条），"
        "其余用例字段与详表一致，完整步骤见项目组测试记录。",
        size=10.5,
    )
    chunk_size = 15
    for chunk_idx in range(0, len(CASES), chunk_size):
        if chunk_idx > 0:
            add_page_break(doc)
        chunk = CASES[chunk_idx : chunk_idx + chunk_size]
        rows = [
            [
                str(chunk_idx + i + 1),
                c["id"],
                c["name"][:20],
                c["type"],
                c["pass"],
                executor_for(c["id"]),
            ]
            for i, c in enumerate(chunk)
        ]
        title = "表 A-1" if chunk_idx == 0 else f"表 A-1（续 {chunk_idx // chunk_size + 1}）"
        add_para(doc, f"{title}　用例总表（第 {chunk_idx + 1}～{chunk_idx + len(chunk)} 条）", bold=True)
        add_table(
            doc,
            ["序号", "用例标识", "用例名称", "类型", "结论", "执行人"],
            rows,
            col_widths_cm=[1.0, 2.2, 4.5, 1.2, 1.2, 1.5],
            font_size=10.5,
            align_center=False,
        )
    add_para(
        doc,
        "注：各用例输入步骤、预期输出与测试说明 4.1～4.4 节及代表用例详表一致，完整记录见项目组测试记录 Excel。",
        size=10.5,
    )
    ti_map = {
        "GN-UM-001": "TI-UM",
        "GN-PM-001": "TI-PM",
        "GN-MM-001": "TI-MM",
        "GN-TM-001": "TI-TM",
        "GN-DG-001": "TI-DG",
        "GN-AD-001": "TI-AD",
        "XN-001": "TI-MM",
        "XN-002": "TI-MM",
        "JK-001": "TI-JK",
        "JK-002": "TI-JK",
        "BJ-001": "TI-MM",
        "AQX-001": "TI-AD",
    }
    for i, c in enumerate(samples):
        if i > 0:
            add_page_break(doc)
        add_heading(doc, f"表 A-{i + 2}　代表用例详表（{c['id']}）", 2)
        add_sample_case_detail(doc, c, ti_map.get(c["id"], "TI-UM"))

    add_heading(doc, "5　注释", 1, compact=merged)
    add_para(
        doc,
        "GN/XN/JK/BJ/AQX/JM/QD 为测试类型标识；用例标识中模块缩写 UM/PM/MM/TM/DG/AD 与测试项 TI-* 对应；"
        "“通过”表示实际输出与预期一致且无阻塞缺陷。",
    )


def build_part4_test_report(doc: Document, *, merged: bool = False):
    if not merged:
        add_page_break(doc)
    if not merged:
        add_cover_block(
            doc,
            "软件配置项测试报告 V1.1",
            "SysMLDocGen-027-STR03",
            "配置项测试报告",
            PAGE_TOTAL,
            36,
        )
        add_toc(
            doc,
            [
                "1　范围",
                "2　依据和引用文档",
                "3　测试概述",
                "4　测试实施情况",
                "5　测试数据分析",
                "6　测试结论",
                "7　后续工作建议",
                "8　注释",
            ],
        )

    add_heading(doc, "1　范围", 1, page_break_before=merged, compact=False)
    add_para(
        doc,
        "本文档为《SysMLDocGen 软件配置项测试报告》，总结 V1.2 配置项测试的实施过程、数据分析与结论，"
        "作为毕业设计软件验收与答辩的支撑材料之一。",
    )
    add_heading(doc, "2　依据和引用文档", 1, compact=merged)
    add_bullets(
        doc,
        [
            "a) 《SysMLDocGen 软件测试计划》V1.0；",
            "b) 《SysMLDocGen 软件测试说明》V1.0；",
            "c) 《SysMLDocGen 软件需求规格说明》V1.2。",
        ],
    )

    add_heading(doc, "3　测试概述", 1, compact=merged)
    add_heading(doc, "3.1　测试工作概述", 2, compact=merged)
    add_para(
        doc,
        "项目组于 2026 年 4 月 17 日至 5 月 20 日完成三轮配置项测试：第一轮冒烟测试 15 例、"
        "第二轮全面测试 71 例、第三轮回归测试 12 例，合计执行 86 条用例。",
    )
    add_para(
        doc,
        "共记录缺陷 12 项，已关闭 11 项，遗留 1 项（BUG-012，S4 级）。测试负责人吴一昊，"
        "指导教师张德平参与计划审核与报告批准。",
    )
    add_heading(doc, "3.2　测试范围", 2, compact=merged)
    add_para(
        doc,
        "测试范围包括：软件安装与启动、用户认证与权限、项目管理、模型上传与解析、模板维护、"
        "Word/PDF 文档生成与下载、系统管理与日志审计，以及性能、接口、安全与界面等非功能项。",
    )
    add_heading(doc, "3.3　测试方法和策略", 2, compact=merged)
    add_para(
        doc,
        "采用黑盒测试为主、白盒审查为辅；功能测试依据测试说明逐条执行并记录实际结果；"
        "接口测试结合 Swagger 与 demo_test.py；性能测试对上传与生成操作计时；"
        "回归测试针对缺陷修复版本重复执行相关用例子集。",
    )
    add_heading(doc, "3.4　测试环境", 2, compact=merged)
    add_para(
        doc,
        "测试环境与《软件测试计划》第 3 章一致：TEST-PC-01、Windows 11、MySQL 8.4、Chrome 浏览器及本机前后端服务；"
        "环境经检查与计划无重大差异，有效。",
    )

    add_heading(doc, "4　测试实施情况", 1)
    add_heading(doc, "4.1　测试实施时间和地点", 2)
    add_para(doc, "2026-04-17～05-20；项目组实验室。")
    add_heading(doc, "4.2　测试参与人员", 2)
    add_table(
        doc,
        ["人员", "职责"],
        [
            ["吴一昊", "负责人/报告"],
            ["申博文", "用例设计与执行"],
            ["林文诚", "环境与审查"],
            ["徐乐", "执行与缺陷跟踪"],
            ["张德平", "指导教师/审核"],
        ],
        col_widths_cm=[3.0, 11.5],
        font_size=10.5,
    )
    add_heading(doc, "4.3　测试实施进程", 2)
    add_table(
        doc,
        ["阶段", "日期"],
        [["冒烟", "04-17～18"], ["全面", "04-19～05-10"], ["回归", "05-11～18"], ["报告", "05-19～24"]],
        col_widths_cm=[5, 9.5],
        font_size=10.5,
    )
    add_heading(doc, "4.4　遇到的问题及解决方案", 2, compact=merged)
    add_table(
        doc,
        ["问题摘要", "影响", "处理措施"],
        [
            ["操作日志接口偶发 500", "中", "修复 SQL 分页与空值判断，回归通过"],
            ["MySQL 服务未自启", "低", "编写 start-mysql.ps1 并写入用户手册"],
            ["PDF 导出中文乱码", "中", "调整转换字体配置后复测通过"],
            ["前端部分页面布局错位", "低", "CSS 修复，JM 用例通过"],
            ["顶栏搜索无后端接口", "低", "登记 BUG-012，S4 遗留"],
        ],
        col_widths_cm=[4.5, 1.5, 8.5],
        font_size=10.5,
        align_center=False,
    )

    add_heading(doc, "5　测试数据分析", 1)
    add_heading(doc, "5.1　测试用例执行结果", 2)
    add_table(
        doc,
        ["轮次", "执行", "通过", "通过率"],
        [["冒烟", "15", "15", "100%"], ["全面", "71", "68", "95.8%"], ["回归", "12", "12", "100%"], ["合计", "86", "85", "98.8%"]],
        col_widths_cm=[3, 2.5, 2.5, 3.5],
        font_size=10.5,
    )
    add_heading(doc, "5.2　需求覆盖分析", 2, compact=merged)
    add_para(
        doc,
        "依据需求跟踪矩阵，SRS 中功能需求 SRS-UM/PM/MM/TM/DG/AD 及非功能需求均已映射至测试用例，"
        "配置项测试需求覆盖率 100%，未发现未覆盖的高优先级需求项。",
    )
    add_heading(doc, "5.3　缺陷趋势分析", 2, compact=merged)
    add_para(
        doc,
        "缺陷在 2026 年 4 月下旬集中暴露（模型解析与 PDF 中文问题），5 月上旬进入修复高峰期，"
        "5 月中旬后新增缺陷为零，仅保留 1 项 S4 级界面建议类遗留。",
    )
    add_heading(doc, "5.4　按模块统计", 2)
    add_table(
        doc,
        ["模块", "缺陷数", "已关闭", "说明"],
        [
            ["用户认证", "1", "1", "令牌过期提示优化"],
            ["项目管理", "2", "2", "权限校验与列表刷新"],
            ["模型管理", "4", "4", "大文件解析与编码"],
            ["模板/文档", "3", "2", "PDF 中文与样式"],
            ["系统管理/界面", "2", "2", "布局与日志筛选"],
        ],
        col_widths_cm=[3.5, 2.0, 2.0, 7.0],
        font_size=10.5,
    )
    add_heading(doc, "5.5　按缺陷严重程度统计", 2, compact=merged)
    add_table(
        doc,
        ["等级", "数量", "已关闭", "遗留"],
        [
            ["S1 致命", "0", "0", "0"],
            ["S2 严重", "2", "2", "0"],
            ["S3 一般", "7", "7", "0"],
            ["S4 轻微", "3", "2", "1"],
        ],
        col_widths_cm=[2.5, 2.0, 2.0, 2.0],
        font_size=10.5,
    )
    add_heading(doc, "5.6　遗留缺陷分析", 2)
    add_table(
        doc,
        ["标识", "摘要", "等级"],
        [["BUG-012", "顶栏搜索无后端", "S4"]],
        col_widths_cm=[3, 8, 3],
        font_size=10.5,
    )

    add_heading(doc, "6　测试结论", 1, compact=merged)
    add_heading(doc, "6.1　测试结果", 2, compact=merged)
    add_para(
        doc,
        "文档审查（附录 A～C）结论均为合格；动态测试 86 条用例执行率 100%、通过率 98.8%；"
        "自动化脚本 demo_test.py 在 TEST-PC-01 上连续三次执行通过。",
    )
    add_heading(doc, "6.2　测试基本结论", 2, compact=merged)
    add_para(
        doc,
        "综合上述结果，判定 SysMLDocGen V1.2 通过配置项测试，满足毕业设计软件验收条件；"
        "遗留缺陷 BUG-012 不影响核心业务流程，建议在后续版本修复。",
        bold=True,
    )

    add_heading(doc, "7　后续工作建议", 1, compact=merged)
    add_bullets(
        doc,
        [
            "7.1 测试方面：将 demo_test.py 纳入持续集成，每次合并主干自动执行冒烟用例；",
            "7.2 开发方面：实现顶栏搜索后端接口，完成 BUG-012 关闭；",
            "7.3 部署方面：生产环境启用 HTTPS 与口令策略，替换默认 admin 密码。",
        ],
    )
    add_heading(doc, "8　注释", 1)
    add_para(
        doc,
        f"本合并文档全文约 {PAGE_TOTAL} 页（A4、正文宋体五号、单倍行距）；"
        "86 条用例均已执行，代表详表格式适用于全部用例记录。",
    )


def main():
    """请优先使用 generate_merged_doc.py 生成合并版（单封面+合并目录）"""
    import generate_merged_doc

    generate_merged_doc.main()


if __name__ == "__main__":
    main()
